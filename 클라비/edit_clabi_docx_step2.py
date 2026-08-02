# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

path = r'C:\Users\NKIA1\hipo-2\IT OKR_KPI-dashboard\클라비\01_클라비_경쟁사분석.docx'
doc = Document(path)


def find_para_by_text(text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise ValueError(f'paragraph not found: {text}')


h6 = find_para_by_text('6. 엔키아 관점의 위협요인')
h8 = find_para_by_text('8. 신용정보 (업데이트 예정)')

# collect all paragraphs from h6 (inclusive) up to h8 (exclusive) in document order
collected = []
el = h6._element
while el is not h8._element:
    collected.append(el)
    el = el.getnext()

# new heading text (renamed, merges 6+7)
h6.text = ''
run = h6.add_run('6. 엔키아 관점의 위협요인 · 협업 가능 영역 · 신규사업 기회')
# heading run picks up style automatically from paragraph style; explicit font name not required

# intro paragraph inserted before h8, then remove old content, then table, then keep trailing separator
intro_p = h8.insert_paragraph_before()
intro_run = intro_p.add_run(
    '엔키아의 실제 경쟁사는 와치텍이며, 클라비는 사업모델이 달라 1:1 경쟁사는 아님. '
    '다만 아래 영역별로 위협요인과 협업 가능 영역, 그리고 엔키아가 클라비의 사업영역을 참고해 '
    '직접 도전해볼 수 있는 신규사업 기회를 함께 정리함.'
)
intro_run.font.name = '맑은 고딕'

rows = [
    ['구분', '위협요인', '협업 가능 영역', '엔키아 신규사업 기회'],
    [
        '공공 AI·클라우드\n통합발주("올인원 패키지")',
        '공공 RFP가 AI+클라우드+운영관리를 하나로 묶어 발주하는 추세. 클라비 같은 통합사업자가 운영관리 영역까지 확장하면 엔키아 폴스타(POLESTAR)가 끼어들 자리가 줄어듦',
        '클라비의 공공 AX 사업에 엔키아를 운영관리 서브파트너로 결합하는 컨소시엄 구성',
        '엔키아도 생성형AI 컨설팅 역량을 갖춰 "AI+운영관리" 통합 제안이 가능한 원사업자로 직접 입찰하는 방안 검토',
    ],
    [
        '네이버클라우드(NCP)\n채널·MSP',
        '클라비는 NCP 최상위 파트너+지분투자까지 받아 공공기관의 NCP 우선검토 시 자동 진입. 엔키아는 전속 CSP 파트너십 부재',
        '클라비를 통해 엔키아 폴스타를 NCP 마켓플레이스에 우회 입점',
        '엔키아가 직접 NCP(또는 타 CSP) 파트너 등급을 취득해 클라우드 MSP 사업에 독자 진출 — 클라비의 MSP(매출 60%) 모델 벤치마크',
    ],
    [
        '생성형AI 에이전트·\nAIOps 결합',
        '클라비의 생성형AI(클라리오) 역량이 향후 IT운영 자동화 영역까지 확장될 가능성',
        '클라비 클라리오(LLM 에이전트) + 엔키아 AIOps 데이터를 결합해 "AI 기반 IT운영 자동화" 공동 제안',
        '엔키아가 자체 운영데이터를 학습시킨 운영 특화 생성형AI 에이전트를 직접 개발해 클라리오와 유사한 자체 브랜드 플랫폼 사업에 진출',
    ],
    [
        '국방·폐쇄망\n특수시장',
        '클라비가 폐쇄망 온프레미스 AI 역량으로 국방시장 진입 — 엔키아와 겹치는 발주처 증가 가능',
        '국방·폐쇄망 사업에서 클라비의 AI 구축 + 엔키아의 운영관리를 결합한 컨소시엄',
        '엔키아도 폐쇄망 온프레미스 AIOps 상품을 직접 개발해 국방·특수망 시장에 독자 진출',
    ],
    [
        'SaaS·구독형\n반복매출 모델',
        '클라비가 구축형에서 구독형·플랫폼 중심으로 전환 중(반복매출 확보) → 공공기관 장기계약을 선점하면 엔키아 신규계약 기회가 줄어들 수 있음',
        '클라비의 구독형 SaaS 플랫폼에 엔키아 모니터링 모듈을 옵션 상품으로 결합 제안',
        '엔키아도 폴스타를 SaaS 구독형 모델로 전환해 반복매출 구조 확보 — 클라비의 SaaS 전환 전략 벤치마크',
    ],
]

tbl = doc.add_table(rows=len(rows), cols=4)
tbl.style = 'Table Grid'
for ri, row_data in enumerate(rows):
    for ci, cell_text in enumerate(row_data):
        cell = tbl.cell(ri, ci)
        cell.text = ''
        p2 = cell.paragraphs[0]
        lines = cell_text.split('\n')
        for li, line in enumerate(lines):
            if li > 0:
                p2.add_run().add_break()
            r2 = p2.add_run(line)
            r2.font.name = '맑은 고딕'
            r2.font.size = Pt(9.5)
            if ri == 0:
                r2.bold = True

# move the table's XML element to right before h8 (after intro_p)
h8_el = h8._element
h8_el.addprevious(tbl._tbl)

# now remove the old collected paragraphs (list items, old separator between 6/7, old heading7, old intro7, etc.)
for el in collected:
    el.getparent().remove(el)

# rename heading 8 -> 7
h8.text = ''
run8 = h8.add_run('7. 신용정보 (업데이트 예정)')

doc.save(path)
print('STEP2 done: sections 6+7 merged into table, section 8 renumbered to 7')
