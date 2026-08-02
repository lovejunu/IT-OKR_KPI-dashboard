# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

path = r'C:\Users\NKIA1\hipo-2\IT OKR_KPI-dashboard\클라비\01_클라비_경쟁사분석.docx'
img_front = r'C:\Users\NKIA1\hipo-2\IT OKR_KPI-dashboard\클라비\안인구대표명함_앞면.jpg'
img_back = r'C:\Users\NKIA1\hipo-2\IT OKR_KPI-dashboard\클라비\안인구대표명함_뒷면.jpg'

doc = Document(path)


def find_para_by_text(text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise ValueError(f'paragraph not found: {text}')


# ---------- 1) Insert business card images into section 4 (before separator leading into "5. 장단점") ----------
anchor = find_para_by_text('5. 장단점')
# the paragraph immediately before this heading is the "───" separator that closes section 4
sep_before_5 = anchor._element.getprevious()
from docx.text.paragraph import Paragraph
sep_before_5_p = Paragraph(sep_before_5, doc)

label_p = sep_before_5_p.insert_paragraph_before()
run = label_p.add_run('[참고] 안인구 대표 명함')
run.bold = True
run.font.name = '맑은 고딕'

front_cap_p = sep_before_5_p.insert_paragraph_before()
front_cap_run = front_cap_p.add_run('앞면')
front_cap_run.font.name = '맑은 고딕'
front_cap_run.font.size = Pt(9)
front_cap_run.italic = True

front_img_p = sep_before_5_p.insert_paragraph_before()
front_img_run = front_img_p.add_run()
front_img_run.add_picture(img_front, width=Inches(4.5))

back_cap_p = sep_before_5_p.insert_paragraph_before()
back_cap_run = back_cap_p.add_run('뒷면')
back_cap_run.font.name = '맑은 고딕'
back_cap_run.font.size = Pt(9)
back_cap_run.italic = True

back_img_p = sep_before_5_p.insert_paragraph_before()
back_img_run = back_img_p.add_run()
back_img_run.add_picture(img_back, width=Inches(4.5))

note_p = sep_before_5_p.insert_paragraph_before()
note_p.paragraph_format.left_indent = Inches(0.3)
note_run = note_p.add_run(
    '※ 위 명함은 약 2년 전(2024년경) 수령한 것으로, 명함에 기재된 주소'
    '(서울특별시 송파구 법원로 114 엠스테이트 A동 903호)는 현재 본사 주소'
    '(서울 송파구 문정동, NH송파농협빌딩 — 위 1번 표 참고)와 다름. 이전한 것으로 추정되며 참고용으로만 활용.'
)
note_run.italic = True
note_run.font.name = '맑은 고딕'
note_run.font.color.rgb = RGBColor(0x33, 0x33, 0x99)

doc.save(path)
print('STEP1 done: images inserted')
