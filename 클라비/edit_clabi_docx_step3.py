# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

path = r'C:\Users\NKIA1\hipo-2\IT OKR_KPI-dashboard\클라비\01_클라비_경쟁사분석.docx'
doc = Document(path)

# ---------- A) widen the "6." threat/collaboration/opportunity table to the current print area ----------
section = doc.sections[0]
usable_inches = section.page_width.inches - section.left_margin.inches - section.right_margin.inches

target_table = None
for t in doc.tables:
    if t.rows[0].cells[0].text.strip() == '구분':
        target_table = t
        break
if target_table is None:
    raise ValueError('target table not found')

target_table.autofit = False
tblPr = target_table._tbl.tblPr
tblLayout = tblPr.find(qn('w:tblLayout'))
if tblLayout is None:
    tblLayout = OxmlElement('w:tblLayout')
    tblPr.append(tblLayout)
tblLayout.set(qn('w:type'), 'fixed')

fracs = [0.13, 0.29, 0.29, 0.29]
widths = [Inches(usable_inches * f) for f in fracs]
for col, w in zip(target_table.columns, widths):
    col.width = w
for row in target_table.rows:
    for idx, cell in enumerate(row.cells):
        cell.width = widths[idx]

# ---------- B) footer with page numbers formatted "- 1 -" ----------
footer = section.footer
footer.is_linked_to_previous = False
if footer.paragraphs:
    fp = footer.paragraphs[0]
    for r in list(fp.runs):
        r._r.getparent().remove(r._r)
else:
    fp = footer.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

run_pre = fp.add_run('- ')
run_pre.font.name = '맑은 고딕'
run_pre.font.size = Pt(10)

run_field = fp.add_run()
run_field.font.name = '맑은 고딕'
run_field.font.size = Pt(10)
fld_begin = OxmlElement('w:fldChar')
fld_begin.set(qn('w:fldCharType'), 'begin')
instr = OxmlElement('w:instrText')
instr.set(qn('xml:space'), 'preserve')
instr.text = 'PAGE'
fld_end = OxmlElement('w:fldChar')
fld_end.set(qn('w:fldCharType'), 'end')
run_field._r.append(fld_begin)
run_field._r.append(instr)
run_field._r.append(fld_end)

run_post = fp.add_run(' -')
run_post.font.name = '맑은 고딕'
run_post.font.size = Pt(10)

# ---------- C) force all text color to black (images untouched) ----------
BLACK = RGBColor(0x00, 0x00, 0x00)


def blacken_paragraph(p):
    for r in p.runs:
        r.font.color.rgb = BLACK


def blacken_table(t):
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                blacken_paragraph(p)
            for nested in cell.tables:
                blacken_table(nested)


# body paragraphs and tables, in document order
from docx.table import Table
from docx.text.paragraph import Paragraph
for child in doc.element.body.iterchildren():
    tag = child.tag.split('}')[-1]
    if tag == 'p':
        blacken_paragraph(Paragraph(child, doc))
    elif tag == 'tbl':
        blacken_table(Table(child, doc))

# header/footer paragraphs across all sections
for sec in doc.sections:
    for hf in (sec.header, sec.footer, sec.first_page_header, sec.first_page_footer,
               sec.even_page_header, sec.even_page_footer):
        for p in hf.paragraphs:
            blacken_paragraph(p)
        for t in hf.tables:
            blacken_table(t)

doc.save(path)
print('STEP3 done: table widened, footer page numbers added, text color set to black')
