from docx import Document
from docx.shared import Pt

path = r'C:\Users\NKIA1\hipo-2\IT OKR_KPI-dashboard\04_신입사원_자기주도학습교재_엔키아이해하기.docx'
doc = Document(path)

target = None
for p in doc.paragraphs:
    if p.text.strip() == '총 약 1,400여 개 고객사 보유':
        target = p
        break
if target is None:
    raise SystemExit('대상 문단을 찾지 못했습니다.')

note = (
    ' (2023년 IR자료 기준. "보유"라는 표현과 현재 거래 중인 고객군을 나열하는 문맥상 현재 시점의 거래 고객사 수로 '
    '추정되나, 원본에 "현재" 또는 "누적"이라는 표현이 명시되어 있지는 않습니다. 역대 누적 구축 고객사 수일 가능성도 '
    '배제할 수 없어, 정확한 정의는 영업본부·경영지원본부에 확인이 필요합니다.)'
)
run = target.add_run(note)
run.italic = True
run.font.name = '맑은 고딕'
run.font.size = Pt(9)

doc.save(path)
print('완료')
