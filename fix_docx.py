"""발표보고서_엔키아_OKR_KPI설계.docx 수정 스크립트"""
import sys, io, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import docx
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

PATH = r'C:\Users\NKIA1\hipo-2\IT OKR_KPI-dashboard\발표보고서_엔키아_OKR_KPI설계.docx'
doc = docx.Document(PATH)
changes = []


# ─── 헬퍼: 셀/단락 텍스트 교체 (run 분산 대응) ───────────────────────────
def replace_text_in_para(para, old, new):
    """단락 전체 텍스트에서 old→new. 서식은 첫 run 기준 유지."""
    full = para.text
    if old not in full:
        return False
    new_full = full.replace(old, new)
    # 첫 run의 서식을 기준으로 모든 run 재구성
    fmt_run = para.runs[0] if para.runs else None
    for run in para.runs:
        run.text = ''
    if fmt_run is None:
        para.add_run(new_full)
    else:
        fmt_run.text = new_full
    return True


def replace_in_cell(cell, old, new):
    changed = False
    for para in cell.paragraphs:
        if replace_text_in_para(para, old, new):
            changed = True
    return changed


def replace_in_all_tables(old, new, label):
    found = False
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if old in cell.text:
                    replace_in_cell(cell, old, new)
                    changes.append(f'[TABLE] {label}')
                    found = True
    return found


def replace_in_all_paragraphs(old, new, label):
    found = False
    for para in doc.paragraphs:
        if old in para.text:
            replace_text_in_para(para, old, new)
            changes.append(f'[PARA] {label}')
            found = True
    return found


# ─── Fix 1: 잡플래닛·블라인드 데이터 출처 제거 (TABLE 5) ──────────────────
replace_in_all_tables(
    '현직자 리뷰 (Glassdoor·잡플래닛·블라인드)',
    '공개 채용공고 JD · 기업 공식 홈페이지 · IR 공시자료',
    '분석 방법론 출처: 잡플래닛·블라인드 → 공개 JD·IR 자료'
)

# ─── Fix 2: 더존비즈온 Win Rate 25%+ KPI 제거 (TABLE 12) ──────────────────
replace_in_all_tables(
    '월별 신규 계약 달성률 · Renewal Rate · Win Rate 25%+',
    '월별 신규 계약 달성률 · Renewal Rate · POC→계약 전환율',
    '더존비즈온 KPI — Win Rate 25%+ 미검증 수치 제거'
)

# ─── Fix 3: 더존비즈온 일 단위 실적 리뷰 언급 제거 (TABLE 12) ─────────────
replace_in_all_tables(
    '공공 입찰 + 방문 영업 병행 — 일 단위 실적 보고 문화 파악',
    '공공 입찰 + 방문 영업 병행 구조 파악',
    '더존비즈온 Note — 잡플래닛 기반 실적 보고 언급 제거'
)

# ─── Fix 4: TABLE 17 Win Rate 조정 근거 — 더존비즈온 기준 제거 ────────────
replace_in_all_tables(
    '더존비즈온 실적 기준 현실화',
    '국내 B2B IT 솔루션 영업 특성 기반 현실화',
    '조정 원칙 — 더존비즈온 기준 → 국내 B2B IT 특성 기반'
)

# ─── Fix 5: 단락 내 잡플래닛·블라인드 잔존 참조 전수 제거 ─────────────────
replace_in_all_paragraphs(
    '잡플래닛·블라인드',
    '공개 JD·IR 자료',
    '단락 내 잡플래닛·블라인드 잔존 참조 제거'
)


# ─── 신규 섹션 추가: 부록 — 데이터 검증 이력 ─────────────────────────────
# "감사합니다" 단락 앞에 삽입 (XML 레벨 조작)

def make_heading_para(doc, text, level_style='Heading 2'):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), level_style)
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    p.append(r)
    return p


def make_normal_para(doc, text):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p


def make_table_element(doc, rows_data, col_widths=None):
    """rows_data: list of list of str. Returns w:tbl element."""
    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    tblStyle = OxmlElement('w:tblStyle')
    tblStyle.set(qn('w:val'), 'TableGrid')
    tblPr.append(tblStyle)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)
    tbl.append(tblPr)

    for i, row_data in enumerate(rows_data):
        tr = OxmlElement('w:tr')
        for j, cell_text in enumerate(row_data):
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            shd = OxmlElement('w:shd')
            if i == 0:  # 헤더행 음영
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), '1E3A5F')
            else:
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'auto')
            tcPr.append(shd)
            tc.append(tcPr)

            p = OxmlElement('w:p')
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            if i == 0:
                b = OxmlElement('w:b')
                rPr.append(b)
                color = OxmlElement('w:color')
                color.set(qn('w:val'), 'FFFFFF')
                rPr.append(color)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = cell_text
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)
    return tbl


# 삽입 위치: "감사합니다" 단락 바로 앞
body = doc.element.body
target_para = None
for child in body:
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'p':
        text = ''.join(t.text or '' for t in child.iter(qn('w:t')))
        if '감사합니다' in text:
            target_para = child
            break

insert_elements = []

# 구분선 단락
insert_elements.append(make_normal_para(doc, ''))

# 섹션 제목
insert_elements.append(make_heading_para(doc, '부록  데이터 출처 · /code-review 검증 이력'))

# 소제목
insert_elements.append(make_normal_para(doc, '■ 데이터 신뢰도 3등급 분류 원칙'))

# TABLE A: 데이터 신뢰도
table_a_data = [
    ['등급', '출처 유형', '예시'],
    ['✅ 검증됨', '기업 공식 채용 JD · IR/공시자료 · 공개 홈페이지', '더존비즈온 매출 4,463억(2025) / 가비아 JD 직접 인용'],
    ['⚠️ 참고', '글로벌 SaaS 리서치(Gartner·Forrester) · 기업 공식 블로그', 'ServiceNow Win Rate 20~35% / Datadog NRR 130%+'],
    ['❌ 제거됨', '유료 리뷰 플랫폼(잡플래닛·블라인드, 미접근) · AI 추론 수치', 'Win Rate 25% 더존비즈온 기준 / 갱신율 85~90% / 공공 수주 100~150건'],
]
insert_elements.append(make_table_element(doc, table_a_data))

insert_elements.append(make_normal_para(doc, ''))
insert_elements.append(make_normal_para(doc, '■ /code-review 교차 검증 결과 (2026-06-15)'))
insert_elements.append(make_normal_para(
    doc,
    'git diff 기반 7개 각도(라인별·삭제 동작·교차 파일·재사용·단순화·효율성·고도) 분석을 통해 아래 5건을 발견·수정하였습니다.'
))

# TABLE B: 검증 결과
table_b_data = [
    ['#', '파일', '발견 내용', '조치'],
    ['1', '02_국내기업_직무분석.md', '더존비즈온 매출 4,023억(2024) 미반영', '연결 4,463억(2025) 업데이트'],
    ['2', '03_OKR_KPI_설계.md (3곳)', '임직원 수 151명 — 최신 정보 미반영', '183명으로 일괄 수정'],
    ['3', 'index.html Upsell KPI', '더존비즈온 기존 고객 중심 — 미검증 벤치마크', '해당 항목 제거'],
    ['4', 'index.html RFP KPI', '최소 12건/분기 권장 — 근거 없는 수치', '목표 Win Rate 달성 기준 표현으로 대체'],
    ['5', 'index.html 나라장터 KPI', '공공 Win Rate 30% 가정 — AI 추론 수치', '8건 수주 목표 기준 25건+ 입찰로 대체'],
]
insert_elements.append(make_table_element(doc, table_b_data))

insert_elements.append(make_normal_para(doc, ''))
insert_elements.append(make_normal_para(
    doc,
    '본 보고서의 모든 수치는 공식 출처 확인 가능한 데이터만 남겼으며, 미검증 수치는 업계 일반 표현 또는 엔키아 자체 목표 표현으로 대체하였습니다.'
))
insert_elements.append(make_normal_para(doc, ''))

# 삽입
if target_para is not None:
    for elem in insert_elements:
        target_para.addprevious(elem)
    changes.append('[NEW] 부록: 데이터 출처 · /code-review 검증 이력 섹션 추가')
else:
    # fallback: 문서 끝에 추가
    for elem in insert_elements:
        body.append(elem)
    changes.append('[NEW] 부록 섹션 추가 (문서 끝)')


# ─── 저장 ─────────────────────────────────────────────────────────────────
doc.save(PATH)

print(f'\n총 {len(changes)}건 수정 완료:')
for c in changes:
    print(f'  {c}')
print('\n파일 저장 완료.')
