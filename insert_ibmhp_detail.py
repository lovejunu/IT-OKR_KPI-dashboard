from docx import Document
from docx.shared import Pt, RGBColor, Inches

path = r'C:\Users\NKIA1\hipo-2\IT OKR_KPI-dashboard\04_신입사원_자기주도학습교재_엔키아이해하기.docx'
doc = Document(path)

paras = doc.paragraphs
target_idx = None
for idx, p in enumerate(paras):
    if p.text.startswith('IBM·HP 등 글로벌 기업과 동일한 포트폴리오'):
        target_idx = idx
        break
if target_idx is None:
    raise SystemExit('대상 문단을 찾지 못했습니다.')

target = paras[target_idx]
anchor = paras[target_idx + 1]

note = target.add_run(' (아래 상세 설명 참고)')
note.italic = True
note.font.name = '맑은 고딕'
note.font.size = Pt(10)


def quote_para(indent=0.3):
    p = anchor.insert_paragraph_before('')
    p.paragraph_format.left_indent = Inches(indent)
    return p


def add_runs(p, segments):
    for text, bold in segments:
        r = p.add_run(text)
        r.italic = True
        r.bold = bold
        r.font.name = '맑은 고딕'
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x99)
    return p


p1 = quote_para()
add_runs(p1, [('왜 "개발+구축+운영+컨설팅을 다 갖췄다"는 것이 특별한 강점인가?', True)])

p2 = quote_para()
add_runs(p2, [('IT 서비스 사업은 보통 아래 4단계로 나뉩니다.', False)])

tbl = doc.add_table(rows=5, cols=3)
tbl.style = 'Table Grid'
data = [
    ('단계', '하는 일', '보통 이 단계만 전문으로 하는 회사 유형'),
    ('① 개발', '솔루션(제품) 자체를 자체 기술로 만드는 것', '제품(라이선스) 벤더'),
    ('② 구축', '만든 제품을 고객사 환경에 설치·연동·커스터마이징 하는 것', 'SI(시스템통합) 업체'),
    ('③ 운영', '구축 이후 시스템을 지속적으로 유지보수·아웃소싱 운영하는 것', 'IT 아웃소싱/MSP 업체'),
    ('④ 컨설팅', '고객의 IT 운영 현황을 진단하고 개선 방향을 제시하는 것', 'IT 컨설팅펌'),
]
for ri, row in enumerate(data):
    for ci, text in enumerate(row):
        cell = tbl.cell(ri, ci)
        cell.text = ''
        rp = cell.paragraphs[0]
        rr = rp.add_run(text)
        rr.font.name = '맑은 고딕'
        rr.font.size = Pt(10)
        if ri == 0:
            rr.bold = True
p2._p.addnext(tbl._tbl)

p3 = quote_para()
add_runs(p3, [(
    '국내 IT 시장에는 이 중 한두 단계만 잘하는 회사가 대부분입니다. 제품만 만드는 회사는 구축·운영을 협력사(SI사)에 맡기고, '
    'SI사는 자체 제품 없이 여러 회사 제품을 조합해 구축만 하며, 아웃소싱 업체는 제품 개발 능력이 없는 경우가 많습니다. '
    '그래서 고객 입장에서는 장애가 나면 "제품 결함인지, 구축 실수인지, 운영 문제인지"를 여러 회사에 따로 물어야 하는 상황이 흔합니다.',
    False,
)])

p4 = quote_para()
add_runs(p4, [
    ('엔키아는 이 4단계를 전부 사내 조직으로 자체 수행합니다.', True),
    (' (1.5 조직 구조와 연결해서 보면 이해가 쉽습니다)', False),
])

bullets = [
    ('개발', ' — NKIA연구소·AI연구소가 POLESTAR/AIOTION을 100% 자체 기술로 개발 (1.6의 특허·SW저작권 보유가 그 증거)'),
    ('구축', ' — 사업본부·IoT사업본부(기술지원팀, 사업수행팀 등)가 고객사 데이터센터·현장에 직접 나가 설치·연동'),
    ('운영', ' — IT Operation Outsourcing 서비스로 구축 이후의 운영·유지보수까지 책임'),
    ('컨설팅', ' — 자체 개발한 진단방법론 NCS-ITOC로 IT 운영 현황을 진단하고 개선안 제시 (2.4 참고)'),
]
for label, rest in bullets:
    pb = quote_para(indent=0.5)
    add_runs(pb, [('• ', False), (label, True), (rest, False)])

p6 = quote_para()
add_runs(p6, [
    ('이런 "풀스택(수직계열화) 모델"은 원래 ', False),
    ('IBM·HP가 수십 년간 글로벌 엔터프라이즈 IT 시장을 지배했던 방식', True),
    (
        '입니다. 하드웨어·소프트웨어를 자체 개발하면서 동시에 컨설팅·구축·운영까지 한 회사가 책임지는 구조인데, '
        '실제로 엔키아 창립 초기(1999년)의 국내 IT운영관리 시장은 IBM·HP 같은 글로벌 기업이 바로 이 모델로 독식하고 있었습니다'
        '(1.2 연혁 "글로벌 기업을 이긴 마켓 리더" 참고). 엔키아는 국내 기술로 이 구조를 그대로 복제해 시장을 뒤집은 회사입니다.',
        False,
    ),
])

p7 = quote_para()
add_runs(p7, [
    ('정리하면', True),
    (
        ' — 고객 입장에서는 책임 소재가 분산되지 않고 한 회사가 끝까지 책임진다는 신뢰를 주고, 엔키아 입장에서는 구축·운영 현장의 '
        '피드백이 바로 연구소의 제품 개선으로 이어지는 선순환 구조를 만들 수 있다는 점에서 경쟁 우위가 됩니다.',
        False,
    ),
])

doc.save(path)
print('삽입 완료')
