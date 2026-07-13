from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
import docx.opc.constants
from docx.oxml import OxmlElement

path = r'C:\Users\NKIA1\hipo-2\IT OKR_KPI-dashboard\04_신입사원_자기주도학습교재_엔키아이해하기.docx'
doc = Document(path)

paras = doc.paragraphs
anchor_idx = None
for idx, p in enumerate(paras):
    if p.text.startswith('해외 레퍼런스'):
        anchor_idx = idx + 1
        break
if anchor_idx is None:
    raise SystemExit('앵커 문단을 찾지 못했습니다.')
anchor = paras[anchor_idx]


def add_hyperlink(paragraph, url, text, size=10):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), '맑은 고딕')
    rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size * 2))
    rPr.append(sz)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def new_para(before=anchor):
    return before.insert_paragraph_before('')


def plain_run(p, text, bold=False, italic=False, size=11, color=None):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = '맑은 고딕'
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return r


def heading3(text):
    p = new_para()
    p.style = doc.styles['Heading 3']
    p.add_run(text)
    return p


def note_para(url_text, url):
    p = new_para()
    plain_run(p, '출처: ', italic=True, size=10)
    add_hyperlink(p, url, url_text, size=10)
    return p


def sub_bold(text):
    p = new_para()
    plain_run(p, text, bold=True)
    return p


def make_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    tbl.style = 'Table Grid'
    for ci, h in enumerate(headers):
        cell = tbl.cell(0, ci)
        cell.text = ''
        rp = cell.paragraphs[0]
        rr = rp.add_run(h)
        rr.font.name = '맑은 고딕'
        rr.font.size = Pt(10)
        rr.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = tbl.cell(ri, ci)
            cell.text = ''
            rp = cell.paragraphs[0]
            if isinstance(val, tuple):
                # list of (text, url_or_None) segments
                for j, (txt, url) in enumerate(val):
                    if j > 0:
                        rp.add_run('  ·  ').font.name = '맑은 고딕'
                    if url:
                        add_hyperlink(rp, url, txt, size=9)
                    else:
                        rr = rp.add_run(txt)
                        rr.font.name = '맑은 고딕'
                        rr.font.size = Pt(9)
            else:
                rr = rp.add_run(val)
                rr.font.name = '맑은 고딕'
                rr.font.size = Pt(9)
    return tbl


# ---------- 2.6 ----------
heading3('2.6 최근 신제품·수주 소식 (엔키아 공식 블로그 기준)')
note_para('엔키아 공식 네이버 블로그', 'https://blog.naver.com/nkia_official')
p = new_para()
plain_run(p, '아래는 신제품·인증·수주·파트너십 소식만 간추린 것으로, 사내문화·동호회·인물 소개 등은 제외했습니다.', italic=True, size=10, color=RGBColor(0x33, 0x33, 0x99))

sub_bold('신제품 · 인증 소식')
product_rows = [
    ('2026.05', "Polestar 10, GS인증 1등급 획득", 'TTA 주관 국가공인 SW품질인증 최고 등급'),
    ('2026.01', "WSS(와이어로프 안전진단), 조달청 혁신제품 지정", '공공조달·시범구매 등 공공 영역 활용 기술로 공식 인정'),
    ('2026.01', "Polestar 10 신제품 출시 기념행사", '전 사원 대상 신제품 발표 및 부서별 전략 공유'),
    ('2026.01', "Polestar ITSM 표준운영관리, GS인증 획득", '공공기관 표준운영절차 의무화(전자정부법 개정) 대응 패키지형 ITSM'),
    ('2025.04', "국토교통부 '스마트 건설 챌린지' 혁신상 수상", 'WSS 와이어로프 안전진단 센서, 105개팀 중 혁신상 수상'),
    ('2024.04', '2023년 유망 SaaS 개발·육성 지원 성과기업 선정', 'POLESTAR ITG를 온프레미스→SaaS로 전환하는 R&D 성과'),
    ('2022.12', '글로벌 SaaS 육성 프로젝트 우수과제 선정(과기정통부 장관상)', 'POLESTAR EMS SaaS화, 국내외 매출 12억 원 이상 달성'),
    ('2022.05', '조달청 우수제품 지정 (POLESTAR EMS v8)', 'AI 머신러닝 장애 사전탐지, 1초 단위 실시간 모니터링 등으로 우수 평가'),
    ('2022.02', 'POLESTAR Automation v3, GS인증 1등급 획득', '국내 유일 1:N IT운영자동화 솔루션'),
    ('2020.04', 'POLESTAR ITG v8, GS인증 1등급 획득', 'ITSM + ITAM 통합 솔루션'),
    ('2020.07', 'AIOTION(아이오션) 와이어로프 센서, KC인증 획득', '소프트웨어뿐 아니라 하드웨어(센서) 영역까지 인증 확보'),
    ('2019.05', 'AIOTION(아이오션), GS 1등급 획득', 'AI 기반 산업용 IoT 플랫폼'),
    ('2018.10', 'POLESTAR XEUS, 신SW상품대상 장관상 수상', '이기종 클라우드 통합운영 플랫폼'),
]
tbl1 = make_table(['날짜', '소식', '한줄 요약'], product_rows)
last_p_before_table1 = doc.paragraphs[[i for i, pp in enumerate(doc.paragraphs) if pp._p is anchor._p][0] - 1]
last_p_before_table1._p.addnext(tbl1._tbl)

sub_bold('수주 · 파트너십 소식')
order_rows = [
    ('2026.01', '일본 서브게이트(Subgate)와 WSS 일본 총대리점 계약 체결', '일본 전역 대상 장기 협력, WSS 글로벌 사업화 본격 진입'),
    ('2025.11', "부산항만공사(BPA) 공동연구 '항만 크레인 와이어로프 상시진단 기술' 성과 공개", '3년간 17건 공동R&D(총 71억 원) 중 하나'),
    ('2025.09', '㈜한화 건설부문, NKIA AI 와이어로프 진단장비(WSS) 도입', '타워크레인 등 건설기계 안전관리에 실전 적용'),
    ('2024.01', '부산항만공사(BPA)-엔키아, 와이어로프 상시 진단시스템 기술개발 사업 착수', '중기부 R&D기금 약 12억 원, 항만 크레인 대상 공동개발'),
    ('2023.01', '엔키아-오케스트로, 클라우드·AIOps 사업 전략적 업무협약 체결', '국내 IT운영관리 1위 × 클라우드 통합관리 1위 기업 간 협력'),
    ('2021.07', '나라장터 종합쇼핑몰에 Polestar ITG v8 등록', '기존 Polestar EMS 8에 이어 ITSM/ITAM 라인도 나라장터 구매 가능'),
]
tbl2 = make_table(['날짜', '소식', '한줄 요약'], order_rows)
sub_bold_para_idx = [i for i, pp in enumerate(doc.paragraphs) if pp._p is anchor._p][0] - 1
doc.paragraphs[sub_bold_para_idx]._p.addnext(tbl2._tbl)

# ---------- 2.7 ----------
heading3('2.7 제품·서비스 소개 영상 (엔키아 공식 유튜브)')
note_para('엔키아 Nkia 유튜브 채널', 'https://www.youtube.com/channel/UCDHM7AP7E4-Gmc3xw3EuJcw')
p = new_para()
plain_run(p, '채널 영상은 대부분 제품·서비스 소개용이며, 아래는 제품군별로 정리한 목록입니다.', italic=True, size=10, color=RGBColor(0x33, 0x33, 0x99))

video_rows = [
    ('Polestar 10 / Lucida AI', [
        ('Polestar 10 소개', 'https://www.youtube.com/watch?v=xFyVOIIaMhI'),
        ('Lucida AI 소개', 'https://www.youtube.com/watch?v=lJKjBVNgRIM'),
        ('공공솔루션마켓 소개', 'https://www.youtube.com/watch?v=-kH9P_ANl60'),
    ]),
    ('POLESTAR EMS', [
        ('EMS8 제품소개', 'https://www.youtube.com/watch?v=MeuJj_6Q65E'),
    ]),
    ('POLESTAR ITSM (기능 데모 5편)', [
        ('계약관리', 'https://www.youtube.com/watch?v=NNkeOfgZrnM'),
        ('점검관리', 'https://www.youtube.com/watch?v=iDpn8q_OvJM'),
        ('SLA', 'https://www.youtube.com/watch?v=WJzaf4Ys8R8'),
        ('IT프로젝트관리', 'https://www.youtube.com/watch?v=5j8bMWjymZc'),
        ('서비스요청', 'https://www.youtube.com/watch?v=5F-Rj1e340c'),
    ]),
    ('POLESTAR Automation', [
        ('일일점검 데모', 'https://www.youtube.com/watch?v=8LJ8Obw45rY'),
        ('윈도우 일괄패치 데모', 'https://www.youtube.com/watch?v=2rLkDJcev70'),
        ('파일 일괄배포 데모', 'https://www.youtube.com/watch?v=RymJbHjZU8E'),
        ('LG CNS 도입사례 1', 'https://www.youtube.com/watch?v=xufs3MiF-zA'),
        ('LG CNS 도입사례 2', 'https://www.youtube.com/watch?v=2a-qN3xoKjk'),
        ('LG CNS 도입사례 3', 'https://www.youtube.com/watch?v=i9BfX7mgVk8'),
    ]),
    ('AIOTION / WSS', [
        ('Monitoring', 'https://www.youtube.com/watch?v=FadELUbBQQM'),
        ('Crane Dashboard', 'https://www.youtube.com/watch?v=cCUQIwfW9z0'),
        ('Rule engine', 'https://www.youtube.com/watch?v=-7Dt3RfcNHo'),
        ('Dashboard', 'https://www.youtube.com/watch?v=Tpq_QBhFUF4'),
        ('Alarm Control', 'https://www.youtube.com/watch?v=9OOkxYek9qA'),
        ('와이어로프 테스터-해상현장', 'https://www.youtube.com/watch?v=ln1nUkeKiS0'),
        ('와이어로프 홍보영상', 'https://www.youtube.com/watch?v=q4d20tp28zA'),
    ]),
    ('공공 SaaS', [
        ('AI Assistance IT 인프라 모니터링(NIA 공공SaaS트랙)', 'https://www.youtube.com/watch?v=VvE_t1shiso'),
    ]),
    ('회사소개', [
        ('회사소개영상', 'https://www.youtube.com/watch?v=n4DWWPZaUSM'),
        ('회사소개영상(Short ver.)', 'https://www.youtube.com/watch?v=J_oDtpAPwVk'),
        ('IT운영관리솔루션 전문기업 (주)엔키아 소개', 'https://www.youtube.com/watch?v=FeUGNcjYzHg'),
    ]),
]
tbl_rows = [(cat, tuple(items)) for cat, items in video_rows]
tbl3 = make_table(['제품/구분', '영상'], tbl_rows)
last_before_tbl3_idx = [i for i, pp in enumerate(doc.paragraphs) if pp._p is anchor._p][0] - 1
doc.paragraphs[last_before_tbl3_idx]._p.addnext(tbl3._tbl)

doc.save(path)
print('삽입 완료')
