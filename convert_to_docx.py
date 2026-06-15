from docx import Document
from docx.shared import Pt, RGBColor, Inches
import re

md_path = r'C:\Users\NKIA1\hipo-2\IT OKR_KPI-dashboard\발표스크립트_엔키아_OKR_KPI_전체_90분.md'
docx_path = r'C:\Users\NKIA1\hipo-2\IT OKR_KPI-dashboard\발표스크립트_엔키아_OKR_KPI_전체_90분.docx'

doc = Document()

style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(11)

with open(md_path, encoding='utf-8') as f:
    lines = f.readlines()

def add_inline(p, content):
    parts = re.split(r'(\*\*[^*]+\*\*)', content)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.name = '맑은 고딕'
        else:
            run = p.add_run(part)
            run.font.name = '맑은 고딕'

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')
    stripped = line.strip()

    if re.match(r'^-{3,}$', stripped):
        p = doc.add_paragraph()
        run = p.add_run('─' * 60)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        i += 1
        continue

    if stripped.startswith('# ') and not stripped.startswith('## '):
        doc.add_heading(stripped[2:], level=1)
        i += 1
        continue

    if stripped.startswith('## '):
        doc.add_heading(stripped[3:], level=2)
        i += 1
        continue

    if stripped.startswith('### '):
        doc.add_heading(stripped[4:], level=3)
        i += 1
        continue

    if stripped.startswith('#### '):
        doc.add_heading(stripped[5:], level=4)
        i += 1
        continue

    if stripped.startswith('|'):
        tbl_lines = []
        j = i
        while j < len(lines) and lines[j].strip().startswith('|'):
            tbl_lines.append(lines[j].strip())
            j += 1
        data_lines = [l for l in tbl_lines if not re.match(r'^\|[-| :]+\|$', l)]
        if data_lines:
            cols = len(data_lines[0].split('|')) - 2
            if cols > 0:
                tbl = doc.add_table(rows=len(data_lines), cols=cols)
                tbl.style = 'Table Grid'
                for ri, row_line in enumerate(data_lines):
                    cells = [c.strip() for c in row_line.split('|')[1:-1]]
                    for ci, cell_text in enumerate(cells[:cols]):
                        cell = tbl.cell(ri, ci)
                        cell.text = ''
                        p2 = cell.paragraphs[0]
                        run = p2.add_run(cell_text)
                        run.font.name = '맑은 고딕'
                        run.font.size = Pt(10)
                        if ri == 0:
                            run.bold = True
        i = j
        continue

    if stripped.startswith('> '):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        content = stripped[2:]
        run = p.add_run(content)
        run.italic = True
        run.font.name = '맑은 고딕'
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x99)
        i += 1
        continue

    if stripped.startswith('- ') or stripped.startswith('* '):
        content = stripped[2:]
        p = doc.add_paragraph(style='List Bullet')
        add_inline(p, content)
        i += 1
        continue

    m = re.match(r'^(\d+)\. (.+)', stripped)
    if m:
        content = m.group(2)
        p = doc.add_paragraph(style='List Number')
        add_inline(p, content)
        i += 1
        continue

    if stripped == '':
        i += 1
        continue

    p = doc.add_paragraph()
    add_inline(p, stripped)
    i += 1

doc.save(docx_path)
print('저장 완료:', docx_path)
