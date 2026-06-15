# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.5)
section.top_margin  = section.bottom_margin = Cm(2.2)

# ── helpers ──
def add_heading(text, level=1, color=(30,64,175)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level==1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15 if level==1 else 12)
    run.font.color.rgb = RGBColor(*color)
    return p

def add_body(text, indent=False, bold=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_tip(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(107,114,128)
    return p

def add_screen_cue(text, lines=None):
    """Orange box: 보고서 화면 큐"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear')
    shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'),'FFF7ED')
    pPr.append(shd)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(154,52,18)
    if lines:
        for line in lines:
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(1)
            p2.paragraph_format.space_after  = Pt(1)
            p2.paragraph_format.left_indent  = Cm(1.2)
            pPr2 = p2._p.get_or_add_pPr()
            shd2 = OxmlElement('w:shd')
            shd2.set(qn('w:val'),'clear')
            shd2.set(qn('w:color'),'auto')
            shd2.set(qn('w:fill'),'FFF7ED')
            pPr2.append(shd2)
            run2 = p2.add_run(line)
            run2.font.size = Pt(9.5)
            run2.font.color.rgb = RGBColor(154,52,18)

def add_demo_line(text):
    """Blue box: 대시보드 데모 큐"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear')
    shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'),'EFF6FF')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(30,64,175)
    return p

def spacer():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  TITLE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('엔키아 OKR·KPI 설계 발표 대본 (5분 버전 v2)')
r.bold = True; r.font.size = Pt(18)
r.font.color.rgb = RGBColor(15,23,42)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(4)
r2 = p2.add_run('발표자: 박준우  |  총 발표 시간: 5분  |  보고서 + 대시보드 라이브 데모')
r2.font.size = Pt(10.5); r2.font.color.rgb = RGBColor(100,116,139)

spacer()

# LEGEND
add_tip('[범례]')
add_tip('  주황 박스 = 보고서(발표보고서_엔키아_OKR_KPI설계.docx)를 화면에 띄우는 구간 및 어느 섹션을 보여줄지 지시')
add_tip('  파란 박스 = 대시보드(index.html)로 전환 후 직접 조작하는 라이브 데모 구간')
add_tip('  기울임 회색 = 발표자 전달 팁 (청중에게 말하지 않는 메모)')

spacer()

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  SECTION 1 — 도입
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
add_heading('1부 | 도입  (0:00 ~ 0:30)')
add_tip('자신감 있게 또렷하게. 30초 안에 끝냅니다.')

add_screen_cue('[화면] 보고서 표지 — "직무 OKR·KPI 설계 보고서 / 글로벌 8개 기업 벤치마킹 × AI 직무분석 기반" 페이지',
    ['-> 보고서를 전체 화면(전체 페이지 보기)으로 표지가 보이도록 띄워 두세요.'])

spacer()
add_body('안녕하세요, 경영지원팀 박준우입니다.')
spacer()
add_body('저는 입사 두 달 차 신입으로, 이번 과제를 통해 글로벌 8개 기업의 직무를 AI로 분석하고,')
add_body('엔키아 영업·연구소·사업수행 세 직무의 OKR·KPI를 직접 설계했습니다.')
spacer()
add_body('오늘은 그 과정과 결과물을, 먼저 보고서로 설명드린 뒤')
add_body('제가 직접 만든 대시보드로 실제 작동하는 모습을 보여드리겠습니다.')

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  SECTION 2 — 분석 대상 + 방법론
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
add_heading('2부 | 무엇을, 어떻게 분석했나  (0:30 ~ 1:30)')
add_tip('"내부 정보가 없으면, 외부 최고 기준을 가져오면 됩니다."')

add_screen_cue('[화면] 보고서 "3부 — 분석 방법론" 섹션 → "3-1. 분석 대상 기업 선정" 표로 스크롤',
    ['-> 글로벌 4개사 / 국내 4개사가 정리된 표가 화면에 보이도록 위치시키세요.'])

spacer()
add_body('핵심 질문은 하나였습니다.')
add_body('"엔키아의 성과를 글로벌 기준으로 측정하면 어떻게 될까?"')
spacer()
add_body('엔키아와 같은 영역에서 일하는 8개 기업을 선정했습니다.', bold=True)
spacer()
add_body('글로벌 4개사:', bold=True)
add_body('  - ServiceNow: IT 서비스 관리(ITSM) 세계 1위. 엔키아 Polestar와 가장 직접 비교됩니다.', indent=True)
add_body('  - Dynatrace: AI 이상탐지 AIOps 글로벌 기준점. Polestar AI 기능 정확도 벤치마크.', indent=True)
add_body('  - Datadog: 실시간 IT 인프라 모니터링 강자.', indent=True)
add_body('  - PagerDuty: 장애 자동화·알림 전문. 복구 시간(MTTR) 기준의 원천.', indent=True)
spacer()
add_body('국내 4개사:', bold=True)
add_body('  - 더존비즈온: 국내 B2B 영업 현실과 공공 입찰 구조를 반영했습니다.', indent=True)
add_body('  - 티맥스소프트·카카오엔터프라이즈: 개발 품질 KPI 기준.', indent=True)
add_body('  - 가비아: IDC 운영 고객 만족도 측정 방식 참고.', indent=True)

add_screen_cue('[화면] 같은 3부 섹션 내 "3-2. 분석 프레임워크" 부분으로 스크롤',
    ['-> MEDDIC / DORA / SLO 세 가지 방법론이 언급된 영역이 보이도록 하세요.'])

spacer()
add_body('분석할 때 3가지 글로벌 방법론을 일관되게 적용했습니다.')
add_body('  - MEDDIC: 영업 파이프라인 정량 평가 방법론', indent=True)
add_body('  - DORA 4 메트릭: Google이 만든 개발팀 성과 측정 기준 4가지', indent=True)
add_body('  - SLO + Error Budget: 운영 신뢰성 수치 관리 방식', indent=True)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  SECTION 3 — 핵심 인사이트
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
add_heading('3부 | 핵심 인사이트 3가지  (1:30 ~ 2:30)')
add_tip('인사이트를 말할 때 보고서 해당 섹션을 화면에서 짚어주면 신뢰도가 올라갑니다.')

add_screen_cue('[화면] 보고서 "4부 — 글로벌 4개사 핵심 인사이트" 섹션으로 이동',
    ['-> ServiceNow 영업 / Dynatrace AI 정확도 내용이 보이는 위치로 맞추세요.'])

spacer()
add_body('8개 기업 분석에서 나온 가장 중요한 인사이트 세 가지입니다.', bold=True)
spacer()

add_body('[인사이트 1] 영업 파이프라인은 목표 매출의 3배가 필요합니다.', bold=True, color=(5,150,105))
add_body('ServiceNow 기준, 10건을 논의해야 3건이 계약됩니다.', indent=True)
add_body('엔키아 올해 목표 230억 기준으로는 690억 수준의 파이프라인이 필요합니다.', indent=True)
spacer()

add_screen_cue('[화면] 같은 4부 내 Dynatrace 섹션으로 스크롤 (AI Precision/Recall 언급 부분)')

spacer()
add_body('[인사이트 2] AI 정확도는 숫자로 공개해야 경쟁력이 됩니다.', bold=True, color=(124,58,237))
add_body('Dynatrace는 AI 이상탐지 Precision 85%, Recall 80%를 공개 목표로 씁니다.', indent=True)
add_body('국내 경쟁사 중 이 수치를 공개한 곳이 없습니다.', indent=True)
add_body('엔키아 Polestar가 이 기준을 공식화하면 차별화 포인트가 됩니다.', indent=True)
spacer()

add_screen_cue('[화면] 보고서 "5부 — 국내 4개사 핵심 인사이트" 섹션 → 더존비즈온 영업 부분',
    ['-> Leading KPI / 선행 지표 관련 내용이 보이는 위치로 스크롤하세요.'])

spacer()
add_body('[인사이트 3] 공공 영업은 Leading KPI(선행 지표)로 선행 관리해야 합니다.', bold=True, color=(220,38,38))
add_body('더존비즈온 분석 결과, 공공기관 계약은 평균 120~180일 소요됩니다.', indent=True)
add_body('매출 결과만 보면 영업 활동이 6개월 뒤에야 보입니다.', indent=True)
add_body('"이번 달 신규 유망 고객 확보 건수"같은 선행 지표를 함께 측정해야 합니다.', indent=True)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  SECTION 4 — OKR·KPI 설계 결과
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
add_heading('4부 | 엔키아 OKR·KPI 설계 결과  (2:30 ~ 3:30)')
add_tip('"글로벌 기준을 그대로 쓰지 않고, 엔키아 현실에 맞게 조정했습니다."')

add_screen_cue('[화면] 보고서 "7부 — OKR·KPI 설계 결과 — 3대 직무" 섹션으로 이동',
    ['-> 영업팀 OKR 표가 화면 위쪽에 보이도록 스크롤하세요.',
     '-> 이후 말하면서 연구소 → 사업수행 순으로 천천히 스크롤합니다.'])

spacer()
add_body('공공기관·금융·민간 고객, Polestar 제품, 183명 규모 — 세 가지 현실을 반영했습니다.')
spacer()
add_body('[영업팀]', bold=True, color=(5,150,105))
add_body('"Polestar로 공공·금융·민간 시장에서 영업 경쟁력을 글로벌 수준으로 끌어올린다"', indent=True)
add_body('KR: 파이프라인 690억 유지 / 신규 계약 논의 분기 30건 / 재계약률 90%+', indent=True)
spacer()
add_body('[연구소]', bold=True, color=(124,58,237))
add_body('"Polestar AIOps의 AI 신뢰성과 개발 속도를 글로벌 경쟁 수준으로 높인다"', indent=True)
add_body('KR: AI Precision 85%+ / 배포 후 30일 고객 장애 0건 / 스프린트 완료율 85%+', indent=True)
spacer()
add_body('[사업수행]', bold=True, color=(220,38,38))
add_body('"고객사 Polestar 운영 환경을 계약 SLA 기준 이상으로 안정적으로 유지한다"', indent=True)
add_body('KR: 시스템 가동률 99.5%+ / P1 장애 복구 2시간 이내 / 구축 일정 준수율 90%+', indent=True)
spacer()
add_body('각 KR마다 측정 공식, 목표값, 글로벌 벤치마크를 포함해 총 45개 KPI를 설계했습니다.')

add_screen_cue('[화면 전환 준비] 보고서는 여기서 마칩니다.',
    ['-> 잠깐 멈추고 "지금부터는 이 설계를 실제로 구현한 대시보드를 보여드리겠습니다."라고 말한 뒤',
     '-> 보고서를 닫고 대시보드(index.html 또는 GitHub Pages URL)를 띄우세요.'])

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  SECTION 5 — LIVE DEMO
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
add_heading('5부 | [LIVE DEMO] 대시보드  (3:30 ~ 4:30)', color=(30,64,175))
add_tip('이 구간이 발표의 하이라이트입니다. 천천히, 화면을 보여주면서 설명하세요.')

add_body('지금까지 말씀드린 설계 결과를 단순한 보고서가 아니라,')
add_body('실제로 작동하는 웹 대시보드로 구현했습니다.')
spacer()
add_body('Claude Code AI로 코드를 작성하고, GitHub 무료 서버에 올렸습니다. 비용 0원입니다.')
spacer()

add_demo_line('[DEMO STEP 1] 대시보드를 전체 화면으로 띄웁니다.')
add_demo_line('   URL: https://lovejunu.github.io/IT-OKR_KPI-dashboard')
add_demo_line('   SCRIPT: "지금 핸드폰으로 이 주소에 접속하시면 바로 보실 수 있습니다."')
add_demo_line('   SCRIPT: "영업·연구소·사업수행 세 개 탭으로 구성되어 있습니다."')
spacer()

add_demo_line('[DEMO STEP 2] 영업 탭 클릭 → OKR 달성 현황과 KPI 목록 보여주기')
add_demo_line('   SCRIPT: "방금 보고서에서 말씀드린 영업팀 OKR 4개가 여기 차트로 보입니다."')
add_demo_line('   SCRIPT: "그 아래 KPI 14개도 글로벌 MEDDIC 기반으로 설계했습니다."')
spacer()

add_demo_line('[DEMO STEP 3] [데모 프리셋] 버튼 클릭')
add_demo_line('   SCRIPT: "현재 이런 실적이 들어왔다고 가정하면 — 버튼 하나로 샘플 데이터가 채워집니다."')
add_demo_line('   -> 달성률 칩(초록/노랑/빨강)과 KPI 평균 달성률 바가 자동 업데이트됩니다.')
add_demo_line('   SCRIPT: "녹색이면 목표 달성(100%+), 노란색이면 80~99%, 빨간색이면 80% 미만입니다."')
spacer()

add_demo_line('[DEMO STEP 4] 입력값 하나를 직접 수정 (예: 신규 Qualified Lead 건수 23 -> 15)')
add_demo_line('   SCRIPT: "제가 지금 실적값을 직접 수정해보겠습니다."')
add_demo_line('   -> 달성률 칩이 즉시 바뀌고, 위의 OKR 달성 현황 차트도 자동 연동됩니다.')
add_demo_line('   SCRIPT: "KPI 하나가 바뀌면 연결된 OKR 달성률도 실시간으로 반영됩니다."')
spacer()

add_demo_line('[DEMO STEP 5] 연구소 탭 -> 사업수행 탭 순서로 전환')
add_demo_line('   SCRIPT: "연구소와 사업수행도 동일한 구조입니다."')
add_demo_line('   SCRIPT: "각 팀장님들의 실제 수치로 업데이트하면, 이 화면이 실제 성과관리 도구로 발전할 수 있습니다."')

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  SECTION 6 — 마무리
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
add_heading('6부 | 마무리  (4:30 ~ 5:00)')
add_tip('마지막 30초. 대시보드 화면을 유지한 채로 마무리하세요.')

add_body('오늘 발표에서 두 가지를 보여드리고 싶었습니다.')
spacer()
add_body('[첫째] AI는 "할 수 없었던 일을 가능하게" 만드는 도구입니다.', bold=True)
add_body('입사 두 달 차 신입이 글로벌 8개 기업을 분석하고 대시보드까지 만든 것은,', indent=True)
add_body('AI 없이는 불가능했습니다.', indent=True)
spacer()
add_body('[둘째] 오늘 결과물은 완성본이 아니라 시작점입니다.', bold=True)
add_body('각 팀의 실제 수치와 피드백이 쌓일수록,', indent=True)
add_body('이 대시보드는 엔키아의 실제 성과관리 체계로 발전할 수 있습니다.', indent=True)
spacer()
add_body('이상으로 발표를 마치겠습니다. 감사합니다.')

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  Q&A
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
add_heading('[참고] 예상 질문 & 답변', color=(71,85,105))

add_body('Q1. 이 KPI를 현업에서 바로 쓸 수 있나요?', bold=True)
add_body('"오늘 내용은 글로벌 기준을 엔키아에 맞게 조정한 초안입니다. 실제 도입하려면 각 팀장님과 목표 수치를 조율하고, 현재 측정 가능한 데이터를 확인하는 과정이 필요합니다. 오늘 발표가 그 대화의 시작점이 되기를 바랍니다."', indent=True)
spacer()
add_body('Q2. OKR과 KPI의 차이가 뭔가요?', bold=True)
add_body('"OKR은 어디로 가야 하는가의 방향, KPI는 지금 얼마나 잘 가고 있는가의 측정값입니다. 한라산 등산에 비유하면 OKR은 목표 봉우리, KPI는 현재 속도·남은 거리입니다."', indent=True)
spacer()
add_body('Q3. AI로 다 만든 건가요?', bold=True)
add_body('"Claude Code AI를 도구로 활용했습니다. 하지만 AI가 알아서 다 한 게 아닙니다. 제가 분석 방향을 정하고, 결과를 검토하고, 엔키아 맥락에 맞게 조정했습니다. 사람의 기획력과 AI의 실행력이 함께한 결과물입니다."', indent=True)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  TIMING TABLE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
add_heading('[참고] 타임라인 요약', color=(71,85,105))

table = doc.add_table(rows=7, cols=4)
table.style = 'Table Grid'
rows_data = [
    ['구간', '내용', '화면', '시간'],
    ['1부', '도입', '보고서 표지', '0:00~0:30'],
    ['2부', '분석 대상 + 방법론', '보고서 3부 (분석 대상 표)', '0:30~1:30'],
    ['3부', '핵심 인사이트 3가지', '보고서 4부·5부 (인사이트)', '1:30~2:30'],
    ['4부', 'OKR·KPI 설계 결과', '보고서 7부 (OKR 표)', '2:30~3:30'],
    ['5부', '[LIVE DEMO] 대시보드 5단계', '대시보드로 전환', '3:30~4:30'],
    ['6부', '마무리', '대시보드 유지', '4:30~5:00'],
]

for r_idx, row_data in enumerate(rows_data):
    for c_idx, val in enumerate(row_data):
        cell = table.rows[r_idx].cells[c_idx]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(val)
        run.bold = (r_idx == 0)
        run.font.size = Pt(9.5)
        if r_idx == 0:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(r'C:\Users\NKIA1\hipo-2\IT OKR_KPI-dashboard\발표대본_엔키아_OKR_KPI설계_5분_260615.docx')
print('DONE')
